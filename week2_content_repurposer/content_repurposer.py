"""AI-powered blog content repurposer.

Usage:
    python week2_content_repurposer/content_repurposer.py https://example.com/blog-post

Scrapes readable article text, asks OpenAI to create platform-specific content,
and saves the result as a formatted DOCX document.
"""

from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openai import OpenAI


@dataclass(frozen=True)
class Settings:
    openai_api_key: str
    model: str
    timeout: int
    max_chars: int

    @classmethod
    def from_env(cls) -> "Settings":
        key = os.getenv("OPENAI_API_KEY")
        if not key:
            raise ValueError("OPENAI_API_KEY is required")
        return cls(
            openai_api_key=key,
            model=os.getenv("OPENAI_MODEL", "gpt-5.6-luna"),
            timeout=int(os.getenv("REQUEST_TIMEOUT", "20")),
            max_chars=int(os.getenv("MAX_ARTICLE_CHARS", "30000")),
        )


def scrape_article(url: str, timeout: int, max_chars: int) -> tuple[str, str]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError("Please provide a valid http(s) blog URL")

    response = requests.get(
        url,
        timeout=timeout,
        headers={"User-Agent": "Mozilla/5.0 (compatible; ContentRepurposer/1.0)"},
    )
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript"]):
        tag.decompose()

    title = soup.title.get_text(" ", strip=True) if soup.title else parsed.netloc
    candidates = soup.find_all(["article", "main"])
    root = max(candidates, key=lambda node: len(node.get_text(" ", strip=True)), default=soup)
    paragraphs = [p.get_text(" ", strip=True) for p in root.find_all("p")]
    text = "\n\n".join(p for p in paragraphs if len(p) >= 40)

    if len(text) < 300:
        text = root.get_text(" ", strip=True)
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        raise ValueError("Could not extract readable article content from the URL")
    return title, text[:max_chars]


def generate_content(client: OpenAI, model: str, title: str, article: str) -> dict[str, object]:
    prompt = f"""Repurpose the following blog post into three original social-media assets.

BLOG TITLE: {title}
BLOG CONTENT:
{article}

Return ONLY valid JSON with exactly these keys:
{{
  "twitter_thread": ["tweet 1", "tweet 2", "..."],
  "linkedin_post": "...",
  "instagram_caption": "..."
}}

Requirements:
- Twitter/X: 6-8 concise posts, each <= 280 characters; the final post should contain a clear takeaway or CTA.
- LinkedIn: professional, useful, skimmable; 150-250 words with a strong opening and CTA.
- Instagram: engaging caption, 80-150 words, with a natural CTA and 5-8 relevant hashtags.
- Preserve factual meaning from the source; do not invent statistics, quotations, or claims.
- Do not mention that AI was used.
"""
    response = client.responses.create(model=model, input=prompt)
    raw = response.output_text.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.IGNORECASE | re.DOTALL).strip()
    data = json.loads(raw)
    if not isinstance(data.get("twitter_thread"), list):
        raise ValueError("AI response did not contain a valid Twitter thread")
    return data


def save_docx(output_dir: str, source_url: str, title: str, content: dict[str, object]) -> Path:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "-", title).strip("-").lower()[:60] or "blog"
    path = output / f"{safe_title}_repurposed_{datetime.now():%Y%m%d_%H%M%S}.docx"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    heading = doc.add_heading("AI Content Repurposing Report", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run("Source: ").bold = True
    p.add_run(source_url)
    p = doc.add_paragraph()
    p.add_run("Original title: ").bold = True
    p.add_run(title)

    doc.add_heading("Twitter / X Thread", level=1)
    for i, tweet in enumerate(content["twitter_thread"], 1):
        doc.add_paragraph(f"{i}. {tweet}", style="List Number")

    doc.add_heading("LinkedIn Post", level=1)
    doc.add_paragraph(str(content["linkedin_post"]))

    doc.add_heading("Instagram Caption", level=1)
    doc.add_paragraph(str(content["instagram_caption"]))

    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.save(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description="Repurpose a blog URL into social-media content.")
    parser.add_argument("url", help="Blog post URL")
    parser.add_argument("--output-dir", default="outputs", help="Directory for generated DOCX")
    args = parser.parse_args()

    settings = Settings.from_env()
    title, article = scrape_article(args.url, settings.timeout, settings.max_chars)
    client = OpenAI(api_key=settings.openai_api_key)
    content = generate_content(client, settings.model, title, article)
    path = save_docx(args.output_dir, args.url, title, content)
    print(f"Saved repurposed content to: {path}")


if __name__ == "__main__":
    main()
